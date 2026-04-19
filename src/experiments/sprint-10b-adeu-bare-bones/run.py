"""Sprint 10B — bare-bones Adeu SDK smoke test.

Substrate proof for Adeu 1.1.0 in the sandbox venv. Generates a synthetic
three-sentence .docx via python-docx, applies three hardcoded edits through
Adeu's Python SDK (no LLM, no agent, no Deep Agents), saves the output, and
inspects the resulting OOXML to confirm track changes are structurally
sound.

Three edits cover the two primary shapes we will need in Sprint 10C:

* Edit 1 — small-span modification (ModifyText "England and Wales" → "New York")
* Edit 2 — insertion via prefix-match anchor (ModifyText where new_text
  begins with target_text; engine detects this as INSERTION in
  ``_apply_single_edit_heuristic`` and writes a single w:ins with no w:del)
* Edit 3 — another small-span modification (ModifyText "good-faith
  negotiation" → "mediation")

Adeu 1.1.0's public SDK does not expose a pure-insertion primitive:
``ModifyText`` requires both ``target_text`` and ``new_text`` as non-empty
strings for the heuristic path. The prefix-match pattern above is the
intended SDK route to produce a pure w:ins (see
``src/adeu/redline/engine.py:739``).

Success criteria (all asserted programmatically below):

1. ``RedlineEngine.process_batch`` returns ``edits_applied == 3`` with no
   ``BatchValidationError``.
2. Output .docx is a valid zip.
3. ``word/document.xml`` parses.
4. Every w:ins and w:del carries ``w:author="Oscar"``.
5. "New York" and "mediation" appear in w:ins text (replacements).
6. "England and Wales" and "good-faith negotiation" appear in w:delText
   (originals preserved, not silently removed).
7. The inserted sentence appears in a standalone w:ins with no adjacent
   w:del (confirming the insertion shape).
"""

from __future__ import annotations

import io
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

from adeu import ModifyText, RedlineEngine

HERE = Path(__file__).parent
INPUT_DOCX = HERE / "input.docx"
OUTPUT_DOCX = HERE / "output.docx"

AUTHOR = "Oscar"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

ORIGINAL_SENTENCES = [
    "This Agreement shall be governed by the laws of England and Wales.",
    "The parties agree to resolve any disputes through good-faith negotiation before commencing litigation.",
    "This Agreement may be amended only in writing signed by both parties.",
]

INSERTION_ANCHOR = "signed by both parties."
INSERTED_SENTENCE = "This Agreement constitutes the entire agreement between the parties."


def write_synthetic_docx(path: Path) -> None:
    doc = Document()
    para = doc.add_paragraph()
    para.add_run(" ".join(ORIGINAL_SENTENCES))
    doc.save(path)


def build_edits() -> list[ModifyText]:
    edit_1 = ModifyText(
        target_text="England and Wales",
        new_text="New York",
    )
    # Insertion: new_text starts with target_text; engine treats the suffix as
    # a pure w:ins anchored right after the target.
    edit_2 = ModifyText(
        target_text=INSERTION_ANCHOR,
        new_text=f"{INSERTION_ANCHOR} {INSERTED_SENTENCE}",
    )
    edit_3 = ModifyText(
        target_text="good-faith negotiation",
        new_text="mediation",
    )
    return [edit_1, edit_2, edit_3]


def apply_edits(input_path: Path, output_path: Path, edits: list[ModifyText]) -> dict:
    stream = io.BytesIO(input_path.read_bytes())
    engine = RedlineEngine(stream, author=AUTHOR)
    result = engine.process_batch(edits)
    output_path.write_bytes(engine.save_to_stream().getvalue())
    return result


def read_document_xml(docx_path: Path) -> bytes:
    with zipfile.ZipFile(docx_path) as zf:
        return zf.read("word/document.xml")


def ooxml_text(element, tag: str) -> str:
    return "".join(t.text or "" for t in element.findall(f".//w:{tag}", NS))


def pretty(element) -> str:
    return etree.tostring(element, pretty_print=True).decode().rstrip()


def inspect_and_report(output_path: Path, apply_result: dict) -> None:
    print(f"=== process_batch result: {apply_result} ===")
    assert apply_result["edits_applied"] == 3, apply_result
    assert apply_result["edits_skipped"] == 0, apply_result

    assert zipfile.is_zipfile(output_path), f"{output_path} is not a valid zip"
    with zipfile.ZipFile(output_path) as zf:
        zip_names = zf.namelist()
    print(f"valid zip ({output_path.stat().st_size} bytes, {len(zip_names)} parts)")
    assert "word/document.xml" in zip_names

    xml_bytes = read_document_xml(output_path)
    root = etree.fromstring(xml_bytes)

    ins_elements = root.findall(".//w:ins", NS)
    del_elements = root.findall(".//w:del", NS)
    print(f"w:ins={len(ins_elements)}  w:del={len(del_elements)}")

    author_attr = f"{{{W_NS}}}author"
    for el in ins_elements + del_elements:
        assert el.get(author_attr) == AUTHOR, (
            f"expected author '{AUTHOR}', got '{el.get(author_attr)}' on {el.tag}"
        )
    print(f"all track-change authors == {AUTHOR!r}")

    ins_texts = [ooxml_text(e, "t") for e in ins_elements]
    del_texts = [ooxml_text(e, "delText") for e in del_elements]

    assert "New York" in "".join(ins_texts), ins_texts
    assert "England and Wales" in "".join(del_texts), del_texts
    assert "mediation" in "".join(ins_texts), ins_texts
    assert "good-faith negotiation" in "".join(del_texts), del_texts
    assert INSERTED_SENTENCE in "".join(ins_texts), ins_texts

    # The insertion edit should produce exactly one w:ins whose text contains
    # the inserted sentence but does not appear in any w:del — insertion, not
    # modification.
    insertion_ins = [e for e, t in zip(ins_elements, ins_texts) if INSERTED_SENTENCE in t]
    assert len(insertion_ins) == 1, insertion_ins
    for del_ in del_elements:
        assert INSERTED_SENTENCE not in ooxml_text(del_, "delText")

    print("\n--- verbatim w:ins fragments ---")
    for i, ins in enumerate(ins_elements, 1):
        print(f"[ins {i}]\n{pretty(ins)}\n")
    print("--- verbatim w:del fragments ---")
    for i, del_ in enumerate(del_elements, 1):
        print(f"[del {i}]\n{pretty(del_)}\n")


def main() -> int:
    print(f"input:  {INPUT_DOCX}")
    print(f"output: {OUTPUT_DOCX}")
    write_synthetic_docx(INPUT_DOCX)
    edits = build_edits()
    apply_result = apply_edits(INPUT_DOCX, OUTPUT_DOCX, edits)
    inspect_and_report(OUTPUT_DOCX, apply_result)
    print("sprint-10b: Adeu bare-bones smoke test passed.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
