"""Sprint 10H — build the synthetic NDA input document.

Identical clause structure to Sprint 10D/10E/10F/10G for cross-sprint
comparability — regenerate locally so the run is self-contained. Only
``OUTPUT_PATH`` changes between sprints.

Run directly to (re)generate the document::

    python src/experiments/sprint-10h/build_input.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

HERE = Path(__file__).parent
OUTPUT_PATH = HERE / "nda-input.docx"

TITLE = "MUTUAL NON-DISCLOSURE AGREEMENT"

PREAMBLE = (
    'This Mutual Non-Disclosure Agreement (this "Agreement") is entered into '
    "as of 1 April 2026 by and between Acme Holdings Limited, a company "
    "incorporated in England and Wales with company number 12345678 whose "
    "registered office is at 1 Example Street, London EC1A 1AA "
    '("Acme"), and Zenith Partners Limited, a company incorporated in '
    "England and Wales with company number 87654321 whose registered office "
    'is at 99 Sample Road, Manchester M1 1AA ("Zenith"). Acme and Zenith '
    'are each referred to as a "Party" and together as the "Parties".'
)

CLAUSES: list[tuple[str, str]] = [
    (
        "1. Definitions",
        (
            'In this Agreement, "Confidential Information" means any '
            "information, whether written, oral, electronic, or in any "
            "other form, disclosed by one Party (the "
            '"Disclosing Party") to the other Party (the "Receiving '
            'Party") in connection with the Purpose, including without '
            "limitation technical, commercial, financial, operational, "
            "and strategic information, whether or not marked or "
            'described as confidential. "Purpose" means the Parties\' '
            "evaluation of a potential commercial collaboration relating "
            "to the supply of professional services by Zenith to Acme."
        ),
    ),
    (
        "2. Confidentiality Obligations",
        (
            "Each Receiving Party shall keep the Disclosing Party's "
            "Confidential Information strictly confidential and shall not "
            "disclose it to any third party without the prior written "
            "consent of the Disclosing Party. The Receiving Party shall "
            "use the Confidential Information solely for the Purpose and "
            "shall take all reasonable steps to protect the Confidential "
            "Information from unauthorised use or disclosure, including "
            "steps at least as stringent as those it takes to protect "
            "its own confidential information of similar importance."
        ),
    ),
    (
        "3. Permitted Disclosures",
        (
            "The Receiving Party may disclose Confidential Information "
            "to its directors, officers, employees, and professional "
            'advisers (together, "Representatives") who have a genuine '
            "need to know it for the Purpose, provided that the "
            "Receiving Party procures that each such Representative is "
            "bound by obligations of confidentiality no less onerous "
            "than those set out in this Agreement. The Receiving Party "
            "shall be liable for any breach of this Agreement by its "
            "Representatives."
        ),
    ),
    (
        "4. Exclusions",
        (
            "The obligations in this Agreement do not apply to "
            "information which: (a) is or becomes publicly available "
            "other than through a breach of this Agreement; (b) was "
            "lawfully in the possession of the Receiving Party without "
            "an obligation of confidence prior to disclosure; (c) is "
            "lawfully obtained from a third party without breach of any "
            "obligation of confidence; or (d) is required to be "
            "disclosed by law, regulation, or order of a court of "
            "competent jurisdiction, provided that the Receiving Party "
            "gives the Disclosing Party prompt written notice where "
            "lawful to do so."
        ),
    ),
    (
        "5. Term",
        (
            "This Agreement shall commence on the date first written "
            "above and shall continue in force for a period of two (2) "
            "years, after which it shall automatically terminate. The "
            "confidentiality obligations in Clause 2 shall survive "
            "termination for a further period of three (3) years."
        ),
    ),
    (
        "6. Return or Destruction of Information",
        (
            "Upon written request of the Disclosing Party, or on "
            "termination of this Agreement, the Receiving Party shall, "
            "at the Disclosing Party's option, promptly return or "
            "destroy all Confidential Information of the Disclosing "
            "Party in its possession or control, together with all "
            "copies, extracts, and notes containing or derived from "
            "such information, and shall certify such return or "
            "destruction in writing. The Receiving Party may retain one "
            "copy of any Confidential Information where required by "
            "applicable law or its reasonable internal record-keeping "
            "policies, provided that such retained copy remains subject "
            "to this Agreement."
        ),
    ),
    (
        "7. Limitation of Liability",
        (
            "Save for liability which cannot be limited or excluded by "
            "applicable law, the total aggregate liability of each "
            "Party to the other under or in connection with this "
            "Agreement, whether in contract, tort (including "
            "negligence), breach of statutory duty, or otherwise, shall "
            "not exceed one hundred thousand pounds sterling (GBP "
            "100,000)."
        ),
    ),
    (
        "8. No Licence",
        (
            "Nothing in this Agreement grants any licence or right in "
            "or to any Confidential Information, patent, copyright, "
            "trademark, or other intellectual property right of either "
            "Party, whether by implication, estoppel, or otherwise."
        ),
    ),
    (
        "9. Governing Law and Dispute Resolution",
        (
            "This Agreement and any dispute or claim arising out of or "
            "in connection with it or its subject matter or formation "
            "(including non-contractual disputes or claims) shall be "
            "governed by and construed in accordance with the laws of "
            "England and Wales. The parties submit to the exclusive "
            "jurisdiction of the courts of England and Wales for the "
            "resolution of all disputes arising out of or in connection "
            "with this Agreement."
        ),
    ),
    (
        "10. General",
        (
            "This Agreement constitutes the entire agreement between "
            "the Parties in relation to its subject matter and "
            "supersedes all prior communications, understandings, and "
            "agreements. No amendment of this Agreement shall be "
            "effective unless in writing and signed by authorised "
            "representatives of both Parties. This Agreement may be "
            "executed in counterparts, each of which shall be deemed an "
            "original, and all of which together shall constitute one "
            "and the same instrument."
        ),
    ),
]

SIGNATURES = [
    "SIGNED for and on behalf of Acme Holdings Limited by an authorised representative.",
    "SIGNED for and on behalf of Zenith Partners Limited by an authorised representative.",
]


def build_document(output_path: Path = OUTPUT_PATH) -> Path:
    """Emit the NDA to ``output_path`` as a .docx and return the path."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(TITLE, level=1)
    doc.add_paragraph(PREAMBLE)

    for heading, body in CLAUSES:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)

    for line in SIGNATURES:
        doc.add_paragraph(line)

    doc.save(output_path)
    return output_path


CLAUSE_9_TEXT = dict(CLAUSES)["9. Governing Law and Dispute Resolution"]


if __name__ == "__main__":
    path = build_document()
    print(f"wrote {path}")
