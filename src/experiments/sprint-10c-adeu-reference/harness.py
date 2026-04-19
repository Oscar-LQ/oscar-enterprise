"""Sprint 10C — shared test harness for the Adeu reference battery.

Small helpers used by the themed test modules under this directory.
All tests are direct-API: no LLM, no agent, no Deep Agents.

Why a shared harness: the battery covers many narrow operations, each
with similar scaffolding (build a synthetic .docx, construct a
``RedlineEngine``, inspect the resulting OOXML). Pulling that into one
module keeps each test file focused on the specific behaviour it
exercises.
"""

from __future__ import annotations

import io
import logging
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Iterable, Optional

import structlog
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# Route Adeu's structlog output through stdlib logging at WARNING level so the
# test reports aren't drowned by INFO-level "Creating new comments part" lines
# (10B surprise #4). Tests that want to inspect log output read stdlib logs.
_LOG_BUFFER: list[logging.LogRecord] = []


class _BufferHandler(logging.Handler):
    def emit(self, record: logging.LogRecord) -> None:
        _LOG_BUFFER.append(record)


logging.basicConfig(level=logging.WARNING, format="%(name)s %(levelname)s %(message)s")
logging.root.addHandler(_BufferHandler())
# Route structlog through stdlib logging so it obeys the WARNING level above.
# Using structlog.stdlib.BoundLogger gives us a full bound-logger surface
# (including .disabled), unlike make_filtering_bound_logger which returns a
# reduced PrintLogger that Adeu's engine.warning/info paths don't always hit
# through the same entry point.
structlog.configure(
    processors=[
        structlog.stdlib.filter_by_level,
        structlog.stdlib.add_log_level,
        structlog.processors.StackInfoRenderer(),
        structlog.processors.format_exc_info,
        structlog.dev.ConsoleRenderer(colors=False),
    ],
    wrapper_class=structlog.stdlib.BoundLogger,
    logger_factory=structlog.stdlib.LoggerFactory(),
    cache_logger_on_first_use=True,
)
# Silence adeu's own loggers at the stdlib layer.
for name in ("adeu", "adeu.redline.engine", "adeu.redline.mapper",
             "adeu.redline.comments", "adeu.ingest", "adeu.markup",
             "adeu.utils.docx"):
    logging.getLogger(name).setLevel(logging.WARNING)


def captured_logs() -> list[logging.LogRecord]:
    """Return a shallow copy of log records captured so far."""
    return list(_LOG_BUFFER)


def clear_logs() -> None:
    _LOG_BUFFER.clear()

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W16DU_NS = "http://schemas.microsoft.com/office/word/2023/wordml/word16du"
NS = {"w": W_NS, "w16du": W16DU_NS}
AUTHOR_ATTR = f"{{{W_NS}}}author"
ID_ATTR = f"{{{W_NS}}}id"

DEFAULT_AUTHOR = "Oscar"


# ----------------------------- synthetic docs -----------------------------


def build_single_paragraph_docx(text: str) -> bytes:
    """Smallest unit: a .docx with one paragraph containing ``text``."""
    doc = Document()
    doc.add_paragraph(text)
    return _to_bytes(doc)


def build_multi_paragraph_docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    return _to_bytes(doc)


def build_formatted_paragraph_docx(segments: list[tuple[str, dict[str, bool]]]) -> bytes:
    """Each (text, {bold, italic}) tuple becomes a separate run."""
    doc = Document()
    p = doc.add_paragraph()
    for text, props in segments:
        run = p.add_run(text)
        if props.get("bold"):
            run.bold = True
        if props.get("italic"):
            run.italic = True
    return _to_bytes(doc)


def _to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------- inspection --------------------------------


def load_xml(docx_bytes: bytes) -> etree._Element:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        xml = zf.read("word/document.xml")
    return etree.fromstring(xml)


def load_zip_parts(docx_bytes: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return sorted(zf.namelist())


def load_part(docx_bytes: bytes, part_name: str) -> Optional[bytes]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        if part_name not in zf.namelist():
            return None
        return zf.read(part_name)


def find_ins(root: etree._Element) -> list[etree._Element]:
    return root.findall(".//w:ins", NS)


def find_del(root: etree._Element) -> list[etree._Element]:
    return root.findall(".//w:del", NS)


def text_of(element: etree._Element, local: str = "t") -> str:
    """Concatenates all ``w:{local}`` descendant text under ``element``."""
    return "".join((t.text or "") for t in element.findall(f".//w:{local}", NS))


def ins_texts(root: etree._Element) -> list[str]:
    return [text_of(e, "t") for e in find_ins(root)]


def del_texts(root: etree._Element) -> list[str]:
    return [text_of(e, "delText") for e in find_del(root)]


def pretty(element: etree._Element) -> str:
    return etree.tostring(element, pretty_print=True).decode().rstrip()


# ----------------------------- test reporter -----------------------------


@dataclass
class TestResult:
    name: str
    passed: bool
    notes: list[str] = field(default_factory=list)
    error: Optional[str] = None


def note(results: list[TestResult], msg: str) -> None:
    if results:
        results[-1].notes.append(msg)


def run_suite(name: str, cases: Iterable[Callable[..., Any]]) -> list[TestResult]:
    """Run each test callable. Each callable must return a TestResult."""
    results: list[TestResult] = []
    for case in cases:
        try:
            r = case()
        except AssertionError as e:
            r = TestResult(name=case.__name__, passed=False, error=f"assertion: {e}")
        except Exception as e:
            r = TestResult(name=case.__name__, passed=False, error=f"{type(e).__name__}: {e}")
        results.append(r)
        status = "PASS" if r.passed else "FAIL"
        print(f"  [{status}] {r.name}")
        for n in r.notes:
            print(f"         · {n}")
        if r.error:
            print(f"         ! {r.error}")
    passed = sum(1 for r in results if r.passed)
    print(f"  ── {name}: {passed}/{len(results)} passed ──")
    return results


def summarise(all_results: list[TestResult]) -> int:
    passed = sum(1 for r in all_results if r.passed)
    total = len(all_results)
    print(f"\n==================== BATTERY: {passed}/{total} passed ====================")
    if passed == total:
        return 0
    print("\nFailures:")
    for r in all_results:
        if not r.passed:
            print(f"  - {r.name}")
            if r.error:
                print(f"      {r.error}")
    return 1
