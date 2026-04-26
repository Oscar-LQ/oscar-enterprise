"""Phase 1 integration test for the MCP port.

Runs the ported state-of-play + counter-propose helpers against
Phase 0.5's synthetic verification input. Three cases:

  Case 1 — DISPATCHER, shared-prefix replacement (GBP 50,000 →
  GBP 100,000). Single-run counterparty w:ins triggers the SURGICAL
  path. Observed shape: "GBP " preserved on counterparty author;
  Acme w:del wraps "50,000" inside counterparty w:ins; sibling
  Acme w:ins contains "100,000". This is MCP's actual runtime
  behaviour on shared-prefix inputs (Phase 0.5 §5's wholesale
  description was the wholesale primitive in isolation).

  Case 2 — DISPATCHER, no-shared-tokens replacement (GBP 50,000 →
  the agreed sum of one hundred thousand pounds sterling).
  Single-run counterparty w:ins; surgical path triggers but with
  zero EQUAL segments collapses to the same nested-del + sibling-
  ins shape as the wholesale primitive. Verifies the brief's
  "MCP target shape" assertion: counterparty's w:ins wraps Acme's
  w:del containing FULL counterparty-original text 'GBP 50,000' +
  sibling Acme w:ins 'the agreed sum...'.

  Case 3 — WHOLESALE primitive called directly (no dispatcher).
  Same input as Case 2 but bypasses the surgical path. Confirms
  the wholesale primitive itself produces the brief's target shape.

If any case fails, surface to Arturs before Phase 2.

Usage:
    python -m src.redline.experiments.sprint-10P.integration_test
"""

import shutil
import sys
from datetime import date
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document

from src.redline.lib.author_config import AuthorConfig
from src.redline.lib.counter_propose_helpers import (
    counter_propose_insertion,
    find_tracked_change_element,
    get_max_revision_id,
)
from src.redline.lib.counter_propose_inplace import counter_propose_on_document
from src.redline.lib.state_of_play import build_state_of_play
from src.redline.lib.timestamp import generate_timestamp

REPO_ROOT = Path(__file__).resolve().parents[4]
INPUT_DOCX = (
    REPO_ROOT
    / "docs/redline/research/sprint-10P-counter-propose-verification-input.docx"
)
OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_CASE_1 = OUTPUT_DIR / "integration-test-output-case1-surgical.docx"
OUTPUT_CASE_2 = OUTPUT_DIR / "integration-test-output-case2-dispatcher-wholesale.docx"
OUTPUT_CASE_3 = OUTPUT_DIR / "integration-test-output-case3-wholesale-direct.docx"

# Brief: also produce the "primary" output the brief names. Symlink Case 2.
OUTPUT_PRIMARY = OUTPUT_DIR / "integration-test-output.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

NO_SHARED_REPLACEMENT = "the agreed sum of one hundred thousand pounds sterling"


def _read_paragraph_xml(docx_path: Path) -> ET.Element:
    """Return the first body paragraph element of the document XML."""
    with ZipFile(docx_path, "r") as z:
        with z.open("word/document.xml") as f:
            root = ET.parse(f).getroot()
    body = root.find(f"{{{W_NS}}}body")
    paragraphs = body.findall(f"{{{W_NS}}}p")
    for p in paragraphs:
        text_runs = []
        for elem in p.iter():
            tag = elem.tag.split("}")[-1]
            if tag in ("t", "delText") and elem.text:
                text_runs.append(elem.text)
        if "cap on liability" in "".join(text_runs):
            return p
    raise AssertionError("target paragraph not found in output")


def _classify_ins(paragraph: ET.Element) -> tuple[list[ET.Element], list[ET.Element]]:
    counterparty: list[ET.Element] = []
    acme: list[ET.Element] = []
    for child in paragraph:
        if child.tag != f"{{{W_NS}}}ins":
            continue
        author = child.get(f"{{{W_NS}}}author") or ""
        if author == "Counterparty Counsel":
            counterparty.append(child)
        elif author == "Acme Counsel":
            acme.append(child)
    return counterparty, acme


def _del_text(elem: ET.Element) -> str:
    return "".join(t.text or "" for t in elem.iter(f"{{{W_NS}}}delText"))


def _ins_text(elem: ET.Element) -> str:
    parts = []
    for t in elem.iter(f"{{{W_NS}}}t"):
        # Skip w:t inside nested w:del (those are technically w:t→w:delText
        # already converted; defensive pass)
        parts.append(t.text or "")
    return "".join(parts)


def _print_shape(case_label: str, paragraph: ET.Element) -> None:
    print(f"  [{case_label}] paragraph children:")
    for child in paragraph:
        tag = child.tag.split("}")[-1]
        author = child.get(f"{{{W_NS}}}author") or ""
        wid = child.get(f"{{{W_NS}}}id") or ""
        text = ""
        if tag == "ins":
            text = "".join(t.text or "" for t in child.iter(f"{{{W_NS}}}t"))
            del_inside = child.findall(f"{{{W_NS}}}del")
            del_summary = ""
            if del_inside:
                del_summary = (
                    f" / nested w:del[author={del_inside[0].get(f'{{{W_NS}}}author')!r}, "
                    f"text={_del_text(del_inside[0])!r}]"
                )
            print(
                f"    <w:ins id={wid} author={author!r} text={text!r}>{del_summary}"
            )
        elif tag == "del":
            print(
                f"    <w:del id={wid} author={author!r} "
                f"text={_del_text(child)!r}>"
            )
        elif tag == "r":
            text = "".join(t.text or "" for t in child.iter(f"{{{W_NS}}}t"))
            if text:
                print(f"    <w:r text={text!r}>")


def case_1_dispatcher_surgical() -> int:
    """Original Phase 0.5 input + shared-prefix replacement."""
    print()
    print("=== CASE 1: dispatcher, shared-prefix replacement ===")
    print(f"  Output: {OUTPUT_CASE_1}")

    state = build_state_of_play(str(INPUT_DOCX))
    target = state.changes[0]

    shutil.copy2(INPUT_DOCX, OUTPUT_CASE_1)
    document = Document(str(OUTPUT_CASE_1))
    config = AuthorConfig(name="Acme Counsel", date_override=date(2026, 4, 26))
    outcomes = counter_propose_on_document(
        document, [(target, "GBP 100,000")], config,
    )
    document.save(str(OUTPUT_CASE_1))

    print(
        f"  outcome: status={outcomes[0].status} method={outcomes[0].method!r}"
    )

    paragraph = _read_paragraph_xml(OUTPUT_CASE_1)
    _print_shape("case1", paragraph)

    cp, ac = _classify_ins(paragraph)
    fails: list[str] = []

    if outcomes[0].method != "surgical":
        fails.append(f"expected method=surgical, got {outcomes[0].method!r}")
    if len(cp) != 1:
        fails.append(f"expected 1 Counterparty w:ins, got {len(cp)}")
    if len(ac) != 1:
        fails.append(f"expected 1 sibling Acme w:ins, got {len(ac)}")

    if not fails:
        cp_ins = cp[0]
        nested_dels = cp_ins.findall(f"{{{W_NS}}}del")
        if len(nested_dels) != 1:
            fails.append(
                f"expected 1 Acme w:del nested in Counterparty w:ins, "
                f"got {len(nested_dels)}"
            )
        else:
            del_t = _del_text(nested_dels[0])
            if del_t != "50,000":
                fails.append(
                    f"surgical: nested w:delText should be '50,000' "
                    f"(differing portion only); got {del_t!r}"
                )
        if _ins_text(ac[0]) != "100,000":
            fails.append(
                f"surgical: sibling Acme w:ins should contain '100,000' "
                f"(differing portion only); got {_ins_text(ac[0])!r}"
            )

    if fails:
        print("  CASE 1 FAILED:")
        for f in fails:
            print(f"    {f}")
        return 1
    print("  CASE 1 PASS — surgical narrow shape (shared prefix preserved)")
    return 0


def case_2_dispatcher_wholesale_equivalent() -> int:
    """Same input + replacement that shares no tokens — surgical path
    collapses to wholesale-equivalent shape."""
    print()
    print("=== CASE 2: dispatcher, no-shared-tokens replacement ===")
    print(f"  Output: {OUTPUT_CASE_2}")

    state = build_state_of_play(str(INPUT_DOCX))
    target = state.changes[0]

    shutil.copy2(INPUT_DOCX, OUTPUT_CASE_2)
    document = Document(str(OUTPUT_CASE_2))
    config = AuthorConfig(name="Acme Counsel", date_override=date(2026, 4, 26))
    outcomes = counter_propose_on_document(
        document, [(target, NO_SHARED_REPLACEMENT)], config,
    )
    document.save(str(OUTPUT_CASE_2))

    print(
        f"  outcome: status={outcomes[0].status} method={outcomes[0].method!r}"
    )

    paragraph = _read_paragraph_xml(OUTPUT_CASE_2)
    _print_shape("case2", paragraph)

    cp, ac = _classify_ins(paragraph)
    fails: list[str] = []

    if len(cp) != 1:
        fails.append(f"expected 1 Counterparty w:ins, got {len(cp)}")
    if len(ac) != 1:
        fails.append(f"expected 1 sibling Acme w:ins, got {len(ac)}")

    if not fails:
        cp_ins = cp[0]
        nested_dels = cp_ins.findall(f"{{{W_NS}}}del")
        if len(nested_dels) != 1:
            fails.append(
                f"expected 1 Acme w:del nested in Counterparty w:ins, "
                f"got {len(nested_dels)}"
            )
        else:
            del_t = _del_text(nested_dels[0])
            if del_t != "GBP 50,000":
                fails.append(
                    f"wholesale-equivalent: nested w:delText should be FULL "
                    f"'GBP 50,000'; got {del_t!r}"
                )
        ac_t = _ins_text(ac[0])
        if ac_t != NO_SHARED_REPLACEMENT:
            fails.append(
                f"wholesale-equivalent: sibling Acme w:ins should contain "
                f"{NO_SHARED_REPLACEMENT!r}; got {ac_t!r}"
            )

    if fails:
        print("  CASE 2 FAILED:")
        for f in fails:
            print(f"    {f}")
        return 1
    print(
        "  CASE 2 PASS — dispatcher produces brief's MCP target shape "
        "(no shared tokens → surgical collapses to wholesale-equivalent)"
    )
    return 0


def case_3_wholesale_direct() -> int:
    """Bypass the dispatcher; call counter_propose_insertion directly."""
    print()
    print("=== CASE 3: wholesale primitive direct call ===")
    print(f"  Output: {OUTPUT_CASE_3}")

    state = build_state_of_play(str(INPUT_DOCX))
    target = state.changes[0]

    shutil.copy2(INPUT_DOCX, OUTPUT_CASE_3)
    document = Document(str(OUTPUT_CASE_3))
    body = document.element.body
    next_id = get_max_revision_id(body) + 1
    timestamp = generate_timestamp(date(2026, 4, 26))

    element = find_tracked_change_element(body, target.ooxml_id)
    counter_propose_insertion(
        element,
        client_author="Acme Counsel",
        timestamp=timestamp,
        replacement_text=NO_SHARED_REPLACEMENT,
        next_id=next_id,
    )
    document.save(str(OUTPUT_CASE_3))

    paragraph = _read_paragraph_xml(OUTPUT_CASE_3)
    _print_shape("case3", paragraph)

    cp, ac = _classify_ins(paragraph)
    fails: list[str] = []

    if len(cp) != 1:
        fails.append(f"expected 1 Counterparty w:ins, got {len(cp)}")
    if len(ac) != 1:
        fails.append(f"expected 1 sibling Acme w:ins, got {len(ac)}")

    if not fails:
        cp_ins = cp[0]
        nested_dels = cp_ins.findall(f"{{{W_NS}}}del")
        if len(nested_dels) != 1:
            fails.append(
                f"expected 1 Acme w:del nested in Counterparty w:ins, "
                f"got {len(nested_dels)}"
            )
        else:
            del_t = _del_text(nested_dels[0])
            if del_t != "GBP 50,000":
                fails.append(
                    f"wholesale: nested w:delText should be FULL "
                    f"'GBP 50,000'; got {del_t!r}"
                )
        ac_t = _ins_text(ac[0])
        if ac_t != NO_SHARED_REPLACEMENT:
            fails.append(
                f"wholesale: sibling Acme w:ins should contain "
                f"{NO_SHARED_REPLACEMENT!r}; got {ac_t!r}"
            )

    if fails:
        print("  CASE 3 FAILED:")
        for f in fails:
            print(f"    {f}")
        return 1
    print(
        "  CASE 3 PASS — wholesale primitive produces brief's MCP target shape"
    )
    return 0


def main() -> int:
    print(f"Repo root: {REPO_ROOT}")
    print(f"Input fixture: {INPUT_DOCX}")
    if not INPUT_DOCX.exists():
        print(f"ERROR: Input fixture not found: {INPUT_DOCX}", file=sys.stderr)
        return 2

    failures = 0
    failures += case_1_dispatcher_surgical()
    failures += case_2_dispatcher_wholesale_equivalent()
    failures += case_3_wholesale_direct()

    # Make Case 2 the "primary" output the brief names
    if OUTPUT_PRIMARY.exists() or OUTPUT_PRIMARY.is_symlink():
        OUTPUT_PRIMARY.unlink()
    shutil.copy2(OUTPUT_CASE_2, OUTPUT_PRIMARY)

    print()
    if failures:
        print(f"INTEGRATION TEST FAILED ({failures} case(s)).")
        return 1
    print("INTEGRATION TEST PASSED — all three cases produce expected shapes.")
    print()
    print(
        "  Case 1 narrow surgical (shared prefix preserved): "
        "matches MCP's actual runtime behaviour on shared-token inputs"
    )
    print(
        "  Case 2 dispatcher wholesale-equivalent (no shared tokens): "
        "matches the brief's MCP target shape via dispatcher"
    )
    print(
        "  Case 3 wholesale primitive direct: "
        "matches the brief's MCP target shape directly"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
