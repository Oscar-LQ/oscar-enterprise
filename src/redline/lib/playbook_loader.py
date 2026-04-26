"""Per-VPS playbook loader with three-level fallback.

Ports the MCP _load_playbook pattern (see
/sandbox/reference-material/claude-plugin-mcp/src/config/loader.py L99-127)
and adapts it for Oscar's per-document-type, Word-doc playbooks. Three
differences from the MCP source:

1. Source is a .docx (not .md) read via python-docx Document(path).
2. Output is rendered to the markdown shape specified in §6 of
   docs/redline/research/sprint-10Q-phase-1-3-design.md so the planner
   prompt's playbook-layer section can consume it verbatim.
3. Per-VPS storage convention is
   src/redline/lib/playbooks/{client}/playbook-{document_type}.docx
   — the {client} component resolves to a single per-VPS directory at
   runtime; the {document_type} component disambiguates playbooks within
   that client. The level-2 glob picks the first alphabetically (matching
   MCP behaviour), suitable for first-pass on a single document type.

Three-level fallback:
- Level 1: explicit_path (path-traversal-checked)
- Level 2: project_dir glob "playbook-*.docx" (alphabetical first wins)
- Level 3: empty string

Empty string is a valid return value. The loader never raises.
"""

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

_MAX_PLAYBOOK_FILE_SIZE = 5_000_000


def load_playbook(
    project_dir: str | None = None,
    explicit_path: str | None = None,
) -> str:
    """Load a playbook .docx and render to the §6 markdown shape.

    Args:
        project_dir: Per-VPS playbook directory. The loader globs
            "playbook-*.docx" in this directory and uses the first match
            alphabetically.
        explicit_path: Optional override path to a specific playbook .docx.
            Takes priority over project_dir.

    Returns:
        Playbook content rendered as markdown per the design-note §6
        contract, or empty string if no playbook is found or readable.
    """
    if explicit_path is not None:
        if _has_path_traversal(explicit_path):
            return ""
        return _read_docx_safe(explicit_path)

    if project_dir is not None:
        project_path = Path(project_dir)
        if project_path.is_dir():
            candidates = sorted(project_path.glob("playbook-*.docx"))
            if candidates:
                return _read_docx_safe(str(candidates[0]))

    return ""


def _read_docx_safe(file_path: str) -> str:
    try:
        path = Path(file_path)
        if not path.is_file():
            return ""
        if path.stat().st_size > _MAX_PLAYBOOK_FILE_SIZE:
            return ""
        doc = Document(str(path))
    except (OSError, PackageNotFoundError, ValueError):
        return ""
    return _render_docx(doc)


def _render_docx(doc) -> str:
    """Render a Document to the §6 markdown shape.

    Heading 1 paragraphs emit as "# {text}\\n".
    Heading 2 paragraphs emit as "## {text}\\n".
    Other (Normal) paragraphs emit as "{text}\\n\\n".

    Heading-then-body produces "# Title\\nBody...\\n\\n"; heading-then-
    heading produces "# Title\\n## Section\\n" (the orphan-body case in
    §6's degeneracy clause).
    """
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "Heading 1":
            parts.append(f"# {text}\n")
        elif style_name == "Heading 2":
            parts.append(f"## {text}\n")
        else:
            parts.append(f"{text}\n\n")
    return "".join(parts)


def _has_path_traversal(file_path: str) -> bool:
    return ".." in Path(file_path).parts
