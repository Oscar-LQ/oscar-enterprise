"""Unit tests for src/redline/lib/playbook_loader.

Covers the three fallback levels (explicit path / project-dir glob /
empty) and the missing-file robustness path. Plus path-traversal
rejection, alphabetical-first selection on multiple matches, and the
§6 paragraph-break convention.
"""

from pathlib import Path

import pytest
from docx import Document

from src.redline.lib.playbook_loader import load_playbook


@pytest.fixture
def playbook_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Test Playbook", level=1)
    doc.add_heading("1. First Category", level=2)
    doc.add_paragraph("Position paragraph for first category.")
    doc.add_paragraph("Fallback paragraph for first category.")
    doc.add_heading("2. Second Category", level=2)
    doc.add_paragraph("Position paragraph for second category.")
    path = tmp_path / "playbook-test.docx"
    doc.save(str(path))
    return path


def test_explicit_path_returns_rendered_content(playbook_docx: Path) -> None:
    result = load_playbook(explicit_path=str(playbook_docx))
    assert "# Test Playbook" in result
    assert "## 1. First Category" in result
    assert "## 2. Second Category" in result
    assert "Position paragraph for first category." in result
    assert "Fallback paragraph for first category." in result


def test_project_dir_glob_finds_playbook(
    playbook_docx: Path, tmp_path: Path
) -> None:
    result = load_playbook(project_dir=str(tmp_path))
    assert "# Test Playbook" in result
    assert "## 1. First Category" in result


def test_empty_when_no_args() -> None:
    assert load_playbook() == ""
    assert load_playbook(project_dir=None, explicit_path=None) == ""


def test_missing_explicit_path_returns_empty(tmp_path: Path) -> None:
    result = load_playbook(explicit_path=str(tmp_path / "no-such-file.docx"))
    assert result == ""


def test_empty_project_dir_returns_empty(tmp_path: Path) -> None:
    empty_dir = tmp_path / "empty-client"
    empty_dir.mkdir()
    assert load_playbook(project_dir=str(empty_dir)) == ""


def test_nonexistent_project_dir_returns_empty(tmp_path: Path) -> None:
    assert load_playbook(project_dir=str(tmp_path / "no-such-dir")) == ""


def test_path_traversal_rejected() -> None:
    assert load_playbook(explicit_path="../etc/passwd") == ""
    assert load_playbook(explicit_path="foo/../bar.docx") == ""


def test_explicit_path_takes_priority_over_project_dir(
    playbook_docx: Path, tmp_path: Path
) -> None:
    other_dir = tmp_path / "other"
    other_dir.mkdir()
    other_doc = Document()
    other_doc.add_heading("Other Playbook", level=1)
    other_path = other_dir / "playbook-other.docx"
    other_doc.save(str(other_path))

    result = load_playbook(
        project_dir=str(other_dir),
        explicit_path=str(playbook_docx),
    )
    assert "# Test Playbook" in result
    assert "Other Playbook" not in result


def test_multiple_playbooks_alphabetical_first_wins(tmp_path: Path) -> None:
    for name, heading in [
        ("playbook-z-last.docx", "Z Playbook"),
        ("playbook-a-first.docx", "A Playbook"),
        ("playbook-m-middle.docx", "M Playbook"),
    ]:
        doc = Document()
        doc.add_heading(heading, level=1)
        doc.save(str(tmp_path / name))

    result = load_playbook(project_dir=str(tmp_path))
    assert "A Playbook" in result
    assert "M Playbook" not in result
    assert "Z Playbook" not in result


def test_paragraph_break_convention(playbook_docx: Path) -> None:
    """Heading followed by single \\n; non-heading body followed by \\n\\n."""
    result = load_playbook(explicit_path=str(playbook_docx))
    assert (
        "## 1. First Category\n"
        "Position paragraph for first category.\n\n"
        "Fallback paragraph for first category.\n\n"
        "## 2. Second Category\n"
    ) in result


def test_heading_to_heading_degenerate_case(tmp_path: Path) -> None:
    """H1 directly followed by H2 should produce '# Title\\n## Section\\n'."""
    doc = Document()
    doc.add_heading("Top Title", level=1)
    doc.add_heading("Top Section", level=2)
    doc.add_paragraph("Body.")
    path = tmp_path / "playbook-headings.docx"
    doc.save(str(path))

    result = load_playbook(explicit_path=str(path))
    assert result.startswith("# Top Title\n## Top Section\nBody.\n\n")


def test_corrupt_docx_returns_empty(tmp_path: Path) -> None:
    bad_path = tmp_path / "playbook-bad.docx"
    bad_path.write_bytes(b"not a real docx file")
    assert load_playbook(explicit_path=str(bad_path)) == ""


def test_oversized_file_returns_empty(tmp_path: Path) -> None:
    """Files above the size limit return empty without attempting to parse."""
    from src.redline.lib.playbook_loader import _MAX_PLAYBOOK_FILE_SIZE

    big_path = tmp_path / "playbook-huge.docx"
    big_path.write_bytes(b"x" * (_MAX_PLAYBOOK_FILE_SIZE + 1))
    assert load_playbook(explicit_path=str(big_path)) == ""


def test_real_phase_1_1_playbook_renders() -> None:
    """End-to-end: the actual Phase 1.1 playbook artefact loads cleanly.

    Sanity check that the loader's rendering matches the §6 spec on the
    real playbook .docx Arturs approved.
    """
    repo_root = Path(__file__).resolve().parents[3]
    playbook_path = (
        repo_root
        / "src"
        / "redline"
        / "lib"
        / "playbooks"
        / "client_placeholder"
        / "playbook-compute-msa.docx"
    )
    if not playbook_path.is_file():
        pytest.skip("Phase 1.1 artefact not present at expected path")

    result = load_playbook(explicit_path=str(playbook_path))
    assert result.startswith("# Customer-Side Compute Capacity MSA Playbook")
    assert "## Preliminary note\n" in result
    assert "## 1. Data residency and sovereignty\n" in result
    assert "## 13. Dispute resolution\n" in result
    assert "## Catch-all guidance\n" in result
    # No bullets, no tables in the playbook source — output should not
    # contain markdown-ish bullet markers either
    assert "\n- " not in result
    assert "\n* " not in result
