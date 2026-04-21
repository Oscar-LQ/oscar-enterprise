"""Narrow markdown → .docx converter for CoSec drafts — Sprint C1.

Deliberate scope (ADR 022): handle only the markdown shapes CoSec
documents actually use. Not pandoc. Not CommonMark-complete. Just enough
for UK private-company secretarial documents as drafted by the C1
drafter agent.

Shapes supported
----------------

* ATX headings: ``# H1`` and ``## H2``.
* Paragraphs: blank-line-separated lines; within a paragraph, line breaks
  are collapsed to spaces.
* Inline emphasis: ``**bold**``. Markdown italics (``*x*`` / ``_x_``) are
  not emphasised — they render as literal asterisks / underscores.
* Unordered lists: ``- item`` and ``* item`` at line start.
* Ordered lists: ``N. item`` at line start.
* Horizontal rule: a line of ``---`` becomes a thematic break (empty
  paragraph; Word renders divider-like).

Not supported (by design)
-------------------------

Tables, blockquotes, code fences, nested lists, links, images, footnotes,
multi-level headings beyond H2, inline italic/code. If the drafter
produces any of these, they render as literal text. A future sprint that
needs any of the above (likely C3 — branding) extends this module
surgically, guided by concrete need rather than speculative completeness.

Why not pandoc
--------------

``apt install pandoc`` was blocked in the C1 sandbox (no sudo / no dpkg
lock); ``pip install pypandoc-binary`` is the same policy problem in a
different wrapper and would add a chunky binary to the venv for no
architectural gain. Keeping the dep surface at zero (python-docx is
already transitively installed via adeu) is the cleaner shape for
single-track Sprint C1. ADR 022 records this decision.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as _DocxDocument

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ULIST_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_OLIST_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
_HR_RE = re.compile(r"^\s*-{3,}\s*$")


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Split ``text`` on ``**bold**`` spans and add them as runs."""
    cursor = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _render_block(doc: _DocxDocument, block: list[str]) -> None:
    """Render a single non-empty block (list of lines) into ``doc``."""
    first = block[0].rstrip()

    if first.startswith("# "):
        doc.add_heading(first[2:].strip(), level=1)
        return
    if first.startswith("## "):
        doc.add_heading(first[3:].strip(), level=2)
        return
    if _HR_RE.match(first):
        doc.add_paragraph()
        return

    if _ULIST_RE.match(first) or _OLIST_RE.match(first):
        for line in block:
            line = line.rstrip()
            u = _ULIST_RE.match(line)
            o = _OLIST_RE.match(line)
            style = None
            content = line
            if u:
                style = "List Bullet"
                content = u.group(1)
            elif o:
                style = "List Number"
                content = o.group(1)
            else:
                # Continuation line in a list item; append to previous.
                prev = doc.paragraphs[-1] if doc.paragraphs else None
                if prev is not None and prev.style is not None and prev.style.name in {"List Bullet", "List Number"}:
                    prev.add_run(" " + line.strip())
                    continue
                style = None
                content = line
            para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            _add_runs_with_bold(para, content)
        return

    # Default: a paragraph. Collapse internal line breaks to spaces; this
    # matches legal-document convention where a paragraph is one run of
    # prose regardless of source wrapping.
    joined = " ".join(line.strip() for line in block)
    para = doc.add_paragraph()
    _add_runs_with_bold(para, joined)


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> Path:
    """Convert ``md_path`` to ``docx_path`` and return the output path.

    Overwrites the output file if it exists. The caller is responsible
    for deciding the output path (naming, collision handling).
    """
    md_text = Path(md_path).read_text(encoding="utf-8")

    # Split into blank-line-separated blocks. Leading/trailing blank lines
    # don't produce empty blocks.
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in md_text.splitlines():
        if line.strip() == "":
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(line)
    if current:
        blocks.append(current)

    doc = Document()
    for block in blocks:
        _render_block(doc, block)

    Path(docx_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return Path(docx_path)
